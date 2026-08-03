#!/usr/bin/env python3
"""Build the Chinese bispectral technical report from its Markdown source.

Requires python-docx and Pillow. Validation figures are rendered from the
final analysis PDFs with pdftoppm, then embedded in a DOCX. The canonical
DOCX-to-PDF conversion and visual QA are performed by the Codex document
rendering workflow outside this script.
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17212B"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F7FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D8DEE7"
GOLD = "C9972D"
PALE_GOLD = "FFF7E3"
RED = "9B1C1C"
WHITE = "FFFFFF"
LINK_BLUE = "0563C1"

# Use one Unicode font in every OOXML font slot. This keeps mixed Chinese,
# Latin and mathematical text stable across Word and LibreOffice exporters.
FONT_LATIN = "Arial Unicode MS"
FONT_CJK = "Arial Unicode MS"
FONT_MONO = "Arial Unicode MS"
PINGFANG_PATH = Path("/System/Library/Fonts/PingFang.ttc")
USABLE_WIDTH_DXA = 9360


def rgb(value):
    return RGBColor.from_string(value)


def set_rpr_font_properties(r_pr, font=FONT_LATIN, east_asia=FONT_CJK):
    """Set all OOXML font slots plus the East Asian language hint."""
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), font)
    r_fonts.set(qn("w:hint"), "eastAsia")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    lang.set(qn("w:bidi"), "zh-CN")
    return r_pr


def set_run_font(run, size=None, bold=None, italic=None, color=INK,
                 font=FONT_LATIN, east_asia=FONT_CJK):
    run.font.name = font
    r_pr = run._element.get_or_add_rPr()
    set_rpr_font_properties(r_pr, font, east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)
    return run


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_bottom_border(paragraph, color=MID_GRAY, size="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, cached, end):
        run._r.append(node)
    set_run_font(run, size=8.5, color=MUTED)


def configure_document_settings(doc):
    """Set Chinese language defaults and request field refresh on open."""
    settings = doc.settings.element
    theme_font_lang = settings.find(qn("w:themeFontLang"))
    if theme_font_lang is None:
        theme_font_lang = OxmlElement("w:themeFontLang")
        settings.append(theme_font_lang)
    theme_font_lang.set(qn("w:val"), "zh-CN")
    theme_font_lang.set(qn("w:eastAsia"), "zh-CN")

    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_widths(rows, caption=None):
    # Table 9 contains compact numeric geometry. Reserve enough width to keep
    # values such as 4096→4096 and 1 / 15500 on one line.
    if caption and caption.startswith("表 9 ") and len(rows[0]) == 6:
        return [1450, 1300, 1050, 1000, 1400, 3160]

    count = len(rows[0])
    maxima = []
    for col in range(count):
        lengths = []
        for row in rows:
            value = re.sub(r"[`*\\]", "", row[col])
            visual = sum(2 if ord(ch) > 127 else 1 for ch in value)
            lengths.append(min(max(visual, 4), 42))
        maxima.append(max(lengths))
    minimum = 850 if count >= 5 else 1050
    raw = [max(minimum, value * 95) for value in maxima]
    scale = USABLE_WIDTH_DXA / sum(raw)
    widths = [max(720, int(round(value * scale))) for value in raw]
    widths[-1] += USABLE_WIDTH_DXA - sum(widths)
    if widths[-1] < 720:
        deficit = 720 - widths[-1]
        widths[-1] = 720
        donor = max(range(len(widths) - 1), key=lambda idx: widths[idx])
        widths[donor] -= deficit
    return widths


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    configure_document_settings(doc)

    doc_defaults = doc.styles.element.find(qn("w:docDefaults"))
    default_r_pr = doc_defaults.find(qn("w:rPrDefault")).find(qn("w:rPr"))
    set_rpr_font_properties(default_r_pr)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    set_rpr_font_properties(normal._element.get_or_add_rPr())
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        set_rpr_font_properties(style._element.get_or_add_rPr())
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        set_rpr_font_properties(style._element.get_or_add_rPr())
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = doc.styles["Caption"]
    caption.font.name = FONT_LATIN
    set_rpr_font_properties(caption._element.get_or_add_rPr())
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    header.paragraph_format.space_after = Pt(3)
    set_run_font(header.add_run("Acycle Bispectral Analysis"), 8.5, True, color=MUTED)
    set_run_font(header.add_run("\t技术报告 · bispectral 分支"), 8.5, color=MUTED)
    add_bottom_border(header)

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    set_run_font(footer.add_run("Acycle Bispectral Analysis · v1.0 · 2026-08-02"), 8.5, color=MUTED)
    set_run_font(footer.add_run("\t第 "), 8.5, color=MUTED)
    add_page_field(footer)
    set_run_font(footer.add_run(" 页"), 8.5, color=MUTED)

    first = section.first_page_footer.paragraphs[0]
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(first.add_run("技术设计、验证与使用说明 · v1.0 · 2026-08-02"), 8.5, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("ACYCLE 方法模块"), 11, True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("Bispectral Analysis"), 30, True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("双谱与双相干分析工具箱技术报告"), 20, True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    set_run_font(p.add_run("算法设计 · 统计推断 · 软件实现 · 验证结果 · 用户手册"), 12.5, color=MUTED)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(22)
    add_bottom_border(rule, GOLD, "18")

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    labels = ["文档版本", "软件分支", "MATLAB", "生成日期"]
    values = ["v1.0", "bispectral", "R2025b", "2026-08-02"]
    for col, value in enumerate(labels):
        cell = table.rows[0].cells[col]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), 9.5, True, color=NAVY)
    for col, value in enumerate(values):
        p = table.rows[1].cells[col].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), 9.5)
    set_table_geometry(table, [2340] * 4)
    doc.add_paragraph()
    add_callout(doc, "交付状态", "功能已集成到 Acycle Timeseries 菜单；两条估计路径、GUI 单一 IAAFT maximum-statistic 正式推断和五类事务性输出均已实现。旧 999-FT 五数据制品仅作历史回归基线；当前 IAAFT 科学数值必须重新计算后才能更新。", PALE_GOLD, GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    set_run_font(p.add_run("实现基线：8eea94423759；报告生成时工作树含未提交 bispectral 功能"), 8.5, color=MUTED)
    doc.add_page_break()


def add_callout(doc, title, body, fill=LIGHTER_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [USABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(title), 10.5, True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    add_inline(p2, body, 10.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


INLINE_PATTERN = re.compile(
    r"(\*\*.*?\*\*|(?<!\w)\*(?!\s).*?(?<!\s)\*(?!\w)|`.*?`|https?://\S+)"
)


def add_hyperlink(paragraph, text, url, size=10.5):
    """Append a styled external hyperlink while preserving run font slots."""
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = paragraph.add_run(text)
    set_run_font(run, size, color=LINK_BLUE)
    run.underline = True
    paragraph._p.remove(run._r)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text, size=10.5):
    text = text.replace("\\|", "|")
    for part in INLINE_PATTERN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_run_font(paragraph.add_run(part[2:-2]), size, True)
        elif part.startswith("*") and part.endswith("*"):
            set_run_font(paragraph.add_run(part[1:-1]), size, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            set_run_font(paragraph.add_run(part[1:-1]), size - 0.5, color=INK,
                         font=FONT_MONO, east_asia=FONT_CJK)
        elif part.startswith(("http://", "https://")):
            add_hyperlink(paragraph, part, part, size)
        else:
            set_run_font(paragraph.add_run(part), size)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    doi_match = re.search(r"https://doi\.org/\S+", text)
    if doi_match and doi_match.start() > 0:
        add_inline(p, text[:doi_match.start()].rstrip())
        set_run_font(p.add_run()).add_break()
        add_inline(p, text[doi_match.start():])
    else:
        add_inline(p, text)
    return p


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    sizes = {1: 16, 2: 13, 3: 11.5}
    colors = {1: BLUE, 2: BLUE, 3: DARK_BLUE}
    set_run_font(paragraph.add_run(text), sizes[level], True, color=colors[level])
    if level == 1 and (text.startswith("10.") or text.startswith("11.") or
                       text.startswith("附录")):
        paragraph.paragraph_format.page_break_before = True


def create_numbering_instance(doc):
    """Create an independent decimal list that explicitly restarts at one."""
    numbering = doc.part.numbering_part.element
    style = doc.styles["List Number"]._element
    style_num_id = style.find(
        f"./{qn('w:pPr')}/{qn('w:numPr')}/{qn('w:numId')}"
    ).get(qn("w:val"))
    base_num = next(
        node for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) == style_num_id
    )
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_num_id = max(
        int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))
    ) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_num_id


def add_list_item(doc, text, numbered=False, num_id=None):
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    if numbered:
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        ilvl = num_pr.get_or_add_ilvl()
        ilvl.set(qn("w:val"), "0")
        number = num_pr.get_or_add_numId()
        number.set(qn("w:val"), str(num_id))
    # Long DOI hyperlinks otherwise tend to leave a lone ``h`` or
    # ``https://doi.`` at the right margin in the reference list.  Keep the
    # citation and its link in one numbered paragraph, but start the DOI on a
    # continuation line so wrapping remains clean and predictable.
    doi_match = re.search(r"https://doi\.org/\S+", text) if numbered else None
    if doi_match:
        add_inline(p, text[:doi_match.start()].rstrip())
        p.add_run().add_break()
        add_inline(p, text[doi_match.start():])
    else:
        add_inline(p, text)


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_together = True
    set_run_font(p.add_run(text), 11.5, color=INK, font="Cambria Math", east_asia=FONT_CJK)


def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.08
    shade_paragraph(p, LIGHT_GRAY)
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, 8.3, color=INK, font=FONT_MONO, east_asia=FONT_CJK)
        if index < len(lines) - 1:
            run.add_break()


def parse_table_line(line):
    sentinel = "__ESCAPED_PIPE__"
    line = line.strip().replace("\\|", sentinel)
    values = [value.strip().replace(sentinel, "|") for value in line.strip("|").split("|")]
    return values


def add_table(doc, rows, caption=None):
    if caption:
        p = doc.add_paragraph(style="Caption")
        p.paragraph_format.keep_with_next = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(p.add_run(caption), 8.5, True, color=MUTED)
    headers = rows[0]
    body = rows[2:]
    widths = table_widths([headers] + body, caption)
    font_size = {2: 8.8, 3: 8.3, 4: 7.8, 5: 7.5}.get(len(headers), 7.2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), font_size, True, color=NAVY)
    for row_index, data in enumerate(body):
        row = table.add_row()
        if row_index % 2:
            for cell in row.cells:
                set_cell_shading(cell, "FAFBFC")
        for index, value in enumerate(data):
            p = row.cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if re.fullmatch(r"[\d.×→<>=%\-+/ ]+", value) else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, value, font_size)
    set_table_geometry(table, widths)


def add_figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    picture = p.add_run().add_picture(str(image_path), width=Inches(6.45))
    doc_pr = picture._inline.docPr
    doc_pr.set("name", caption.split("。", 1)[0])
    doc_pr.set("title", caption.split("。", 1)[0])
    doc_pr.set("descr", caption)
    cp = doc.add_paragraph(style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # A slightly smaller size keeps unusually long captions on one line and
    # avoids leaving a final one- or two-character Chinese orphan.
    caption_size = 8.2 if len(caption) > 75 else 8.5
    add_inline(cp, caption, caption_size)


def render_validation_figures(validation_dir, target_dir, pdftoppm):
    files = {
        "FIG_SYNTHETIC": "synthetic-qpc-bispectral-1",
        "FIG_AR1": "ar1-rednoise-0.7-bispectral-1",
        "FIG_NEWARK": "newark2km-rsp0.85-bispectral-1",
        "FIG_SITE1262": "site1262-xrf-fe-bispectral-1",
        "FIG_LR04": "lr04-benthic-d18o-bispectral-1",
    }
    mapping = {}
    for token, result_name in files.items():
        current = validation_dir / result_name / f"{result_name}.pdf"
        generic = validation_dir / result_name / "figure.pdf"
        legacy = validation_dir / f"{result_name}.pdf"
        if current.is_file():
            source = current
        elif generic.is_file():
            source = generic
        elif legacy.is_file():
            source = legacy
        else:
            raise FileNotFoundError(
                f"Expected current {current}, generic {generic}, or legacy {legacy}"
            )
        stem = target_dir / token.lower()
        subprocess.run(
            [pdftoppm, "-f", "1", "-l", "1", "-singlefile", "-png", "-r", "110", str(source), str(stem)],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        output = stem.with_suffix(".png")
        if not output.is_file() or output.stat().st_size == 0:
            raise RuntimeError(f"Failed to render {source}")
        mapping[token] = output
    return mapping


def create_workflow(path):
    image = Image.new("RGB", (1800, 500), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype(str(PINGFANG_PATH), 34)
    small = ImageFont.truetype(str(PINGFANG_PATH), 22)
    boxes = [
        ("主列表选数据", "两列数值文件"),
        ("参数窗口", "计算参数\n即时显示参数"),
        ("严格输入检查", "不清洗·不插值\n不做整段去趋势"),
        ("双谱估计", "WOSA 或频率平滑"),
        ("显著性", "IAAFT max-stat\n全图 FWER"),
        ("绘图与保存", "PDF · FIG · MAT\nCSV · JSON"),
    ]
    x0, y0, box_w, box_h, gap = 35, 145, 245, 185, 52
    for index, (title, detail) in enumerate(boxes):
        x = x0 + index * (box_w + gap)
        fill = "#E8EEF5" if index < 3 else ("#FFF7E3" if index < 5 else "#F2F4F7")
        outline = "#2E74B5" if index < 3 else ("#C9972D" if index < 5 else "#5B6573")
        draw.rounded_rectangle((x, y0, x + box_w, y0 + box_h), radius=18,
                               fill=fill, outline=outline, width=4)
        title_box = draw.textbbox((0, 0), title, font=font)
        draw.text((x + (box_w - title_box[2]) / 2, y0 + 38), title, font=font, fill="#17365D")
        detail_box = draw.multiline_textbbox((0, 0), detail, font=small,
                                             spacing=5, align="center")
        detail_width = detail_box[2] - detail_box[0]
        detail_height = detail_box[3] - detail_box[1]
        detail_x = x + (box_w - detail_width) / 2 - detail_box[0]
        detail_y = y0 + 100 + (62 - detail_height) / 2 - detail_box[1]
        draw.multiline_text((detail_x, detail_y), detail, font=small,
                            spacing=5, align="center", fill="#4D5866")
        if index < len(boxes) - 1:
            start = (x + box_w + 8, y0 + box_h / 2)
            end = (x + box_w + gap - 8, y0 + box_h / 2)
            draw.line((start, end), fill="#7F8C99", width=5)
            draw.polygon([(end[0], end[1]), (end[0] - 16, end[1] - 10),
                          (end[0] - 16, end[1] + 10)], fill="#7F8C99")
    title = "Acycle 双谱分析的数据流与软件调用链"
    box = draw.textbbox((0, 0), title, font=font)
    draw.text(((1800 - box[2]) / 2, 48), title, font=font, fill="#17365D")
    image.save(path, dpi=(180, 180))


def parse_markdown(doc, source, figures):
    lines = source.read_text(encoding="utf-8").splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("## 摘要"))
    lines = lines[start:]
    index = 0
    pending_caption = None
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("```"):
            code = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index].rstrip())
                index += 1
            add_code(doc, code)
            index += 1
            continue
        image_match = re.fullmatch(r"!\[(.*?)\]\(\{\{(.*?)\}\}\)", stripped)
        if image_match:
            add_figure(doc, figures[image_match.group(2)], image_match.group(1))
            index += 1
            continue
        if stripped.startswith("$$") and stripped.endswith("$$"):
            add_equation(doc, stripped[2:-2])
            index += 1
            continue
        if stripped.startswith("> "):
            body = stripped[2:].strip()
            title_match = re.match(r"\*\*(.*?)\*\*\s*(.*)", body)
            if title_match:
                title, text = title_match.groups()
            else:
                title, text = "重要说明", body
            accent = RED if ("解释" in title or "原则" in title) else (GOLD if "限制" in title else BLUE)
            fill = PALE_GOLD if accent in (RED, GOLD) else LIGHTER_BLUE
            add_callout(doc, title, text, fill, accent)
            index += 1
            continue
        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:], 3)
            index += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], 2)
            index += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], 1)
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith("表 "):
            pending_caption = stripped
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(parse_table_line(lines[index]))
                index += 1
            add_table(doc, table_lines, pending_caption)
            pending_caption = None
            continue
        if re.match(r"^\d+\.\s+", stripped):
            num_id = create_numbering_instance(doc)
            while index < len(lines):
                numbered = re.match(r"^\d+\.\s+", lines[index].strip())
                if not numbered:
                    break
                add_list_item(doc, numbered.string[numbered.end():],
                              numbered=True, num_id=num_id)
                index += 1
            continue
        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:])
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if (candidate.startswith(("#", "|", ">", "- ", "```", "$$", "![", "表 ")) or
                    re.match(r"^\d+\.\s+", candidate)):
                break
            paragraph_lines.append(candidate)
            index += 1
        add_paragraph(doc, " ".join(paragraph_lines).replace("  ", " "))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--validation-dir", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--pdftoppm", default=shutil.which("pdftoppm"))
    args = parser.parse_args()
    if not args.pdftoppm:
        raise RuntimeError("pdftoppm is required to embed validation figures")
    args.output.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="acycle-bispectral-report-") as temp_name:
        temp = Path(temp_name)
        figures = render_validation_figures(args.validation_dir, temp, args.pdftoppm)
        workflow = temp / "workflow.png"
        create_workflow(workflow)
        figures["FIG_WORKFLOW"] = workflow

        doc = Document()
        configure_document(doc)
        add_cover(doc)
        parse_markdown(doc, args.source, figures)

        props = doc.core_properties
        props.title = "Acycle Bispectral Analysis 双谱与双相干分析工具箱技术报告"
        props.subject = "算法设计、统计推断、软件实现、验证结果与用户手册"
        props.author = "Acycle project"
        props.keywords = "Acycle, bispectrum, bicoherence, paleoclimate, MATLAB, surrogate"
        props.comments = "Generated from TECHNICAL_REPORT_zh_CN.md"
        props.version = "1.0"
        props.created = datetime(2026, 8, 2)
        props.modified = datetime(2026, 8, 2)
        doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
