#!/usr/bin/env python3
"""Build publication-quality Word reports for the COCO validation suite.

The report builder is deliberately independent of MATLAB and Excel.  It reads
the portable audit files emitted by the publication runner:

* a root ``manifest.json`` (or a CSV manifest);
* one ``case_manifest.json`` per data set;
* key/value ``parameters.csv`` and ``summary.csv`` files;
* a plain-text conclusion; and
* PNG figures.

By default, the script writes one ``case_report.docx`` in every case directory
and ``COCO_publication_validation_report.docx`` at the result root.  All writes
are atomic.  Missing optional figures are reported on stderr and skipped; a
missing manifest or a missing required Python dependency is a hard error.

Example
-------
python3 code/test/build_coco_publication_reports.py /path/to/result/manifest.json

The document geometry is A4 portrait with 15 mm margins.  Full-width figures
are capped at 180 mm and paired half-column figures at 88 mm each.
"""

from __future__ import annotations

import argparse
import csv
import datetime as _datetime
import json
import os
from pathlib import Path
import re
import sys
import tempfile
from typing import Any, Iterable, Mapping, Sequence

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor

    _DOCX_IMPORT_ERROR: Exception | None = None
except Exception as exc:  # pragma: no cover - exercised on machines without it
    Document = None  # type: ignore[assignment]
    _DOCX_IMPORT_ERROR = exc

try:
    from PIL import Image
except Exception:  # pragma: no cover - Pillow is optional
    Image = None  # type: ignore[assignment]


SCRIPT_VERSION = "1.0.0"
A4_WIDTH_MM = 210.0
A4_HEIGHT_MM = 297.0
MARGIN_MM = 15.0
FULL_FIGURE_WIDTH_MM = 180.0
HALF_FIGURE_WIDTH_MM = 88.0
MAX_FULL_FIGURE_HEIGHT_MM = 220.0
MAX_HALF_FIGURE_HEIGHT_MM = 165.0


STANDARD_METHODS = (
    (
        "Confirmatory Blocked cvCOCO",
        "The depth series is divided at its midpoint into segments A and B. "
        "In each direction, the training half is used to locate an adaptive "
        "training sedimentation rate and estimate four orbital-group weights "
        "(long eccentricity, short eccentricity, obliquity, and precession). "
        "Those weights are then frozen before the independent held-out half is "
        "searched over the preregistered sedimentation-rate grid. The analysis "
        "is repeated in the reverse direction.",
    ),
    (
        "Confirmatory statistic and null",
        "The primary decision statistic is p_robust = max(p_A, p_B), with "
        "p_robust < 0.05 required for a robust bidirectional detection. The "
        "symmetric full-pipeline p_sym is secondary and cannot rescue a failed "
        "direction. Under the fitted joint null, the separately regularized "
        "halves are independent stationary Gaussian AR(1) series with their "
        "own estimated lag-one coefficients. Every Monte Carlo replicate "
        "repeats training, target freezing, held-out validation, and the full "
        "sedimentation-rate search. Plus-one Monte Carlo p-values and 95% "
        "Wilson intervals quantify simulation precision.",
    ),
    (
        "Exploratory Adaptive COCO",
        "Adaptive COCO learns target amplitudes from the full record and tests "
        "that same record. Its rate-search-corrected global p-value is the "
        "exploratory inferential result (recommended threshold 0.05); the local "
        "p-value at a rate is descriptive, with 0.01 used only as a diagnostic "
        "guide. Adaptive COCO is useful for discovery and rate localization but "
        "is not held-out confirmation.",
    ),
    (
        "Sampling and spectral preprocessing",
        "Depth coordinates are checked for finite values, ordered, and made "
        "unique. If a record or held-out segment is not regularly spaced, it is "
        "interpolated on a grid defined by the median depth increment and the "
        "operation is recorded in the parameters and run log. Periodograms are "
        "compared only on the common native temporal-frequency grid within the "
        "configured maximum temporal frequency. Orbital periods are grouped "
        "noncoherently; participation counts disclose whether a result uses all "
        "nine requested periods or only a frequency-resolved subset.",
    ),
)


STANDARD_LIMITATIONS = (
    "A positive COCO association is spectral evidence consistent with orbital "
    "forcing; it is not an unconditional proof of causation.",
    "The Monte Carlo calibration is conditional on the selected stationary "
    "Gaussian AR(1) null, its plug-in lag-one estimates, and the implemented "
    "preprocessing and target-construction algorithm.",
    "The directional p_A and p_B values are full-pipeline results under the "
    "joint null that both halves are AR(1); they are not conditional tests that "
    "allow the opposite training half to be an arbitrary fixed signal.",
    "Adjacent midpoint halves need not be geologically independent and can "
    "differ in duration, preservation, noise structure, or sedimentation rate.",
    "Rate-grid limits, rate step, frequency cutoff, detrending, padding, AR(1) "
    "option, Monte Carlo size, and random seed should be fixed before inspecting "
    "the significance curves. Unreported post-hoc tuning invalidates the stated "
    "error control.",
    "When fewer than nine periods participate, a positive result applies only "
    "to the resolved partial orbital target and must not be described as "
    "confirmation of the complete nine-period target.",
    "Interpolation can regularize the frequency axis but cannot recover lost "
    "information; long gaps or strongly uneven support require a sensitivity "
    "analysis beyond the interpolation log.",
    "Adaptive COCO uses the same data to estimate target amplitudes and evaluate "
    "association. Its global p-value is algorithm-calibrated but its scientific "
    "role remains exploratory.",
    "A result close to 0.05 is Monte Carlo-uncertain when its 95% binomial "
    "interval crosses 0.05; additional simulations should be run under a "
    "predeclared precision rule.",
    "Negative controls assess the implemented workflow under particular noise "
    "realizations; one or a few controls do not establish the complete false-"
    "positive rate of the method.",
)


class ReportInputError(RuntimeError):
    """Raised when a report input cannot be interpreted safely."""


class WarningLog:
    """Collect nonfatal report-generation warnings and mirror them to stderr."""

    def __init__(self) -> None:
        self.messages: list[str] = []

    def add(self, message: str) -> None:
        message = str(message).strip()
        if message and message not in self.messages:
            self.messages.append(message)
            print(f"WARNING: {message}", file=sys.stderr)


def _require_docx() -> None:
    if Document is None:
        detail = f" ({_DOCX_IMPORT_ERROR})" if _DOCX_IMPORT_ERROR else ""
        raise ReportInputError(
            "python-docx is required to build Word reports"
            f"{detail}. Install it with: python3 -m pip install --user python-docx"
        )


def _text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, float):
        if value != value:
            return "NaN"
        return f"{value:.10g}"
    if isinstance(value, (list, tuple)):
        return ", ".join(_text(item) for item in value)
    if isinstance(value, Mapping):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return str(value).strip()


def _as_list(value: Any) -> list[Any]:
    if value is None or value == "":
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, tuple):
        return list(value)
    return [value]


def _read_json(path: Path) -> dict[str, Any]:
    try:
        with path.open("r", encoding="utf-8-sig") as handle:
            value = json.load(handle)
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        raise ReportInputError(f"Cannot read JSON manifest {path}: {exc}") from exc
    if not isinstance(value, dict):
        raise ReportInputError(f"Manifest {path} must contain one JSON object.")
    return value


def _read_text(path: Path, warnings: WarningLog) -> str:
    if not path.is_file():
        warnings.add(f"Optional conclusion file is unavailable: {path}")
        return ""
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding).strip()
        except UnicodeError:
            continue
        except OSError as exc:
            warnings.add(f"Cannot read optional text file {path}: {exc}")
            return ""
    warnings.add(f"Cannot decode optional text file: {path}")
    return ""


def _read_csv_records(path: Path, warnings: WarningLog) -> list[dict[str, str]]:
    if not path.is_file():
        warnings.add(f"Optional CSV file is unavailable: {path}")
        return []
    try:
        raw = path.read_bytes()
    except OSError as exc:
        warnings.add(f"Cannot read optional CSV file {path}: {exc}")
        return []

    encoding = "utf-8-sig"
    try:
        decoded = raw.decode(encoding)
    except UnicodeError:
        encoding = "gb18030"
        try:
            decoded = raw.decode(encoding)
        except UnicodeError as exc:
            warnings.add(f"Cannot decode optional CSV file {path}: {exc}")
            return []
    sample = decoded[:8192]

    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
    except csv.Error:
        dialect = csv.excel
    try:
        with path.open("r", encoding=encoding, newline="") as handle:
            rows = [
                {str(key or "").strip(): _text(value) for key, value in row.items()}
                for row in csv.DictReader(handle, dialect=dialect)
            ]
    except (OSError, UnicodeError, csv.Error) as exc:
        warnings.add(f"Cannot parse optional CSV file {path}: {exc}")
        return []
    return [row for row in rows if any(value for value in row.values())]


def _rows_as_metrics(records: Sequence[Mapping[str, Any]]) -> list[dict[str, str]]:
    """Normalize common MATLAB/Python CSV layouts to Method/Metric/Value."""

    if not records:
        return []
    headers = {key.casefold(): key for row in records for key in row}
    metric_key = next(
        (headers[name] for name in ("metric", "parameter", "name", "field") if name in headers),
        None,
    )
    value_key = next(
        (headers[name] for name in ("value", "setting", "result") if name in headers),
        None,
    )
    method_key = next(
        (headers[name] for name in ("method", "analysis", "mode") if name in headers),
        None,
    )
    if metric_key and value_key:
        return [
            {
                "Method": _text(row.get(method_key, "")) if method_key else "",
                "Metric": _text(row.get(metric_key, "")),
                "Value": _text(row.get(value_key, "")),
            }
            for row in records
            if _text(row.get(metric_key, ""))
        ]

    # The publication runner writes one wide row per method/role, for
    # example method, analysis_role, classification, p_robust, p_sym, ....
    # Expand that layout while retaining enough provenance to distinguish
    # the primary run from background-sensitivity rows.
    if method_key:
        role_key = next(
            (headers[name] for name in ("analysis_role", "role") if name in headers),
            None,
        )
        normalized: list[dict[str, str]] = []
        for row in records:
            method = _text(row.get(method_key, ""))
            role = _text(row.get(role_key, "")) if role_key else ""
            method_label = f"{method} [{role}]" if role else method
            for key, value in row.items():
                if key in {method_key, role_key}:
                    continue
                normalized.append(
                    {
                        "Method": method_label,
                        "Metric": _text(key),
                        "Value": _text(value),
                    }
                )
        return normalized

    if len(records) == 1:
        row = records[0]
        return [
            {"Method": "", "Metric": _text(key), "Value": _text(value)}
            for key, value in row.items()
            if _text(key)
        ]

    normalized: list[dict[str, str]] = []
    for index, row in enumerate(records, 1):
        nonempty = [(key, value) for key, value in row.items() if _text(value)]
        if len(nonempty) == 2:
            normalized.append(
                {
                    "Method": "",
                    "Metric": _text(nonempty[0][1]),
                    "Value": _text(nonempty[1][1]),
                }
            )
        else:
            for key, value in nonempty:
                normalized.append(
                    {
                        "Method": "",
                        "Metric": f"Row {index}: {key}",
                        "Value": _text(value),
                    }
                )
    return normalized


def _normalized_key(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", "", _text(value).casefold())


def _find_metric(
    rows: Sequence[Mapping[str, str]],
    include: Sequence[str],
    method_include: Sequence[str] = (),
) -> str:
    include_norm = tuple(_normalized_key(item) for item in include)
    method_norm = tuple(_normalized_key(item) for item in method_include)
    for row in rows:
        metric = _normalized_key(row.get("Metric", ""))
        method = _normalized_key(row.get("Method", ""))
        if all(token in metric for token in include_norm) and (
            not method_norm or any(token in method for token in method_norm)
        ):
            value = _text(row.get("Value", ""))
            if value:
                return value
    return ""


def _resolve_file(root: Path, case_dir: Path, value: Any, fallback: str = "") -> Path:
    raw = _text(value, fallback)
    if not raw:
        return case_dir / "__missing_optional_file__"
    path = Path(os.path.expanduser(raw))
    if path.is_absolute():
        return path
    candidate = case_dir / path
    if candidate.exists() or not (root / path).exists():
        return candidate
    return root / path


def _manifest_from_csv(path: Path, warnings: WarningLog) -> dict[str, Any]:
    records = _read_csv_records(path, warnings)
    cases: list[dict[str, Any]] = []
    for index, row in enumerate(records, 1):
        lowered = {key.casefold(): value for key, value in row.items()}
        case: dict[str, Any] = dict(row)
        case.setdefault("id", lowered.get("id", f"case_{index:02d}"))
        figure_value = lowered.get("figures", lowered.get("figure_paths", ""))
        if figure_value:
            try:
                figures = json.loads(figure_value)
            except json.JSONDecodeError:
                figures = [item.strip() for item in figure_value.split(";") if item.strip()]
            case["figures"] = figures
        cases.append(case)
    return {"cases": cases, "report_title": "COCO publication validation report"}


def _find_root_manifest(
    source: Path, warnings: WarningLog
) -> tuple[Path, dict[str, Any]]:
    source = source.expanduser().resolve()
    if source.is_file():
        if source.suffix.casefold() == ".json":
            return source.parent, _read_json(source)
        if source.suffix.casefold() in {".csv", ".tsv"}:
            return source.parent, _manifest_from_csv(source, warnings)
        raise ReportInputError(f"Unsupported manifest type: {source}")
    if not source.is_dir():
        raise ReportInputError(f"Result root or manifest does not exist: {source}")
    for filename in ("manifest.json", "run_manifest.json", "cases.json"):
        candidate = source / filename
        if candidate.is_file():
            return source, _read_json(candidate)
    for filename in ("manifest.csv", "run_manifest.csv", "cases.csv"):
        candidate = source / filename
        if candidate.is_file():
            return source, _manifest_from_csv(candidate, warnings)
    discovered = sorted(source.glob("*/case_manifest.json"))
    if discovered:
        return source, {"cases": [{"case_manifest": str(path)} for path in discovered]}
    raise ReportInputError(
        f"No manifest found at {source}. Expected manifest.json or */case_manifest.json."
    )


def _case_directory(root: Path, entry: Mapping[str, Any], index: int) -> Path:
    raw = entry.get("case_dir") or entry.get("directory") or entry.get("path")
    if raw:
        path = Path(os.path.expanduser(_text(raw)))
        return path if path.is_absolute() else root / path
    case_manifest = entry.get("case_manifest")
    if case_manifest:
        path = Path(os.path.expanduser(_text(case_manifest)))
        path = path if path.is_absolute() else root / path
        return path.parent
    case_id = _text(entry.get("id"), f"case_{index:02d}")
    return root / case_id


def _merge_case_manifest(root: Path, entry: Mapping[str, Any], index: int) -> dict[str, Any]:
    case_dir = _case_directory(root, entry, index)
    manifest_path_raw = entry.get("case_manifest")
    if manifest_path_raw:
        manifest_path = Path(os.path.expanduser(_text(manifest_path_raw)))
        manifest_path = manifest_path if manifest_path.is_absolute() else root / manifest_path
    else:
        manifest_path = case_dir / "case_manifest.json"
    merged = dict(entry)
    if manifest_path.is_file():
        merged.update(_read_json(manifest_path))
        # The physical location of case_manifest.json is authoritative for
        # relative files unless the root explicitly supplies a case_dir.
        if not entry.get("case_dir"):
            case_dir = manifest_path.parent
    merged["_case_dir"] = case_dir.resolve()
    merged["_case_manifest"] = manifest_path.resolve()
    merged.setdefault("id", f"case_{index:02d}")
    merged.setdefault("title", merged["id"])
    return merged


def _figure_caption_from_name(path: Path) -> str:
    words = re.sub(r"[_-]+", " ", path.stem).strip()
    return words[:1].upper() + words[1:] + "." if words else "COCO result."


def _normalize_figure(
    root: Path,
    case_dir: Path,
    value: Any,
    warnings: WarningLog,
) -> dict[str, Any] | None:
    item: dict[str, Any]
    if isinstance(value, Mapping):
        item = dict(value)
        raw_path = item.get("path") or item.get("file") or item.get("filename")
    else:
        item = {}
        raw_path = value
    if not raw_path:
        warnings.add(f"A figure entry for {case_dir.name} has no path and was skipped.")
        return None
    path = _resolve_file(root, case_dir, raw_path)
    if path.suffix.casefold() not in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
        png = path.with_suffix(".png")
        if png.is_file():
            path = png
        else:
            warnings.add(
                f"Word cannot embed optional figure {path.name}; no PNG counterpart was found."
            )
            return None
    if not path.is_file():
        warnings.add(f"Optional figure is unavailable and was skipped: {path}")
        return None
    if Image is not None:
        try:
            with Image.open(path) as image:
                image.verify()
        except Exception as exc:
            warnings.add(f"Optional figure is unreadable and was skipped: {path} ({exc})")
            return None
    width = _text(item.get("width", "full")).casefold()
    width_class = "half" if width in {"half", "half-column", "88", "88mm"} else "full"
    method = _text(item.get("method", ""))
    if not method:
        filename = path.name.casefold()
        method = "Adaptive COCO" if "adaptive" in filename else (
            "Blocked cvCOCO" if "cvcoco" in filename or "cv_coco" in filename else "COCO"
        )
    return {
        "path": path,
        "caption": _text(item.get("caption"), _figure_caption_from_name(path)),
        "width": width_class,
        "method": method,
        "order": item.get("order", 0),
        "_order_explicit": "order" in item,
    }


def _load_case(root: Path, entry: Mapping[str, Any], index: int, warnings: WarningLog) -> dict[str, Any]:
    warning_start = len(warnings.messages)
    case = _merge_case_manifest(root, entry, index)
    case_dir = Path(case["_case_dir"])
    case["_root"] = root
    parameter_path = _resolve_file(
        root, case_dir, case.get("parameters_csv"), "parameters.csv"
    )
    summary_path = _resolve_file(root, case_dir, case.get("summary_csv"), "summary.csv")
    conclusion_path = _resolve_file(
        root, case_dir, case.get("conclusion_txt"), "conclusion.txt"
    )
    case["_parameter_path"] = parameter_path
    case["_summary_path"] = summary_path
    case["_conclusion_path"] = conclusion_path
    case["_parameters"] = _rows_as_metrics(_read_csv_records(parameter_path, warnings))
    case["_summary"] = _rows_as_metrics(_read_csv_records(summary_path, warnings))
    case["_conclusion"] = _read_text(conclusion_path, warnings)

    raw_figures = _as_list(case.get("figures"))
    if not raw_figures:
        raw_figures = [str(path) for path in sorted((case_dir / "figures").glob("*.png"))]
    figures: list[dict[str, Any]] = []
    for sequence, raw in enumerate(raw_figures):
        figure = _normalize_figure(root, case_dir, raw, warnings)
        if figure is not None:
            figure["_sequence"] = sequence
            figures.append(figure)

    def sort_key(figure: Mapping[str, Any]) -> tuple[float, int]:
        try:
            order = float(figure.get("order", 0))
        except (TypeError, ValueError):
            order = 0.0
        return order, int(figure.get("_sequence", 0))

    if any(bool(figure.get("_order_explicit")) for figure in figures):
        figures.sort(key=sort_key)
    case["_figures"] = figures
    case["_warnings"] = warnings.messages[warning_start:]
    return case


def load_suite(source: Path, warnings: WarningLog) -> tuple[Path, dict[str, Any], list[dict[str, Any]]]:
    root, manifest = _find_root_manifest(source, warnings)
    entries = manifest.get("cases", [])
    if isinstance(entries, Mapping):
        entries = [dict(value, id=key) if isinstance(value, Mapping) else {"id": key, "case_dir": value}
                   for key, value in entries.items()]
    if not isinstance(entries, list) or not entries:
        raise ReportInputError("The suite manifest must define a nonempty cases array.")
    cases: list[dict[str, Any]] = []
    for index, entry in enumerate(entries, 1):
        if isinstance(entry, str):
            entry = {"case_dir": entry, "id": Path(entry).name}
        if not isinstance(entry, Mapping):
            raise ReportInputError(f"Case entry {index} is not a JSON object or path.")
        cases.append(_load_case(root, entry, index, warnings))
    return root, manifest, cases


def _set_font(run: Any, name: str, size_pt: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _set_cell_margins(cell: Any, top: int = 55, start: int = 70, bottom: int = 55, end: int = 70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def _configure_document(doc: Any, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Mm(A4_WIDTH_MM)
    section.page_height = Mm(A4_HEIGHT_MM)
    section.top_margin = Mm(MARGIN_MM)
    section.bottom_margin = Mm(MARGIN_MM)
    section.left_margin = Mm(MARGIN_MM)
    section.right_margin = Mm(MARGIN_MM)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size, color in (
        ("Title", 17, "17365D"),
        ("Subtitle", 11, "44546A"),
        ("Heading 1", 14, "17365D"),
        ("Heading 2", 12, "1F4E79"),
        ("Heading 3", 10.5, "1F4E79"),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8 if "Heading" in style_name else 0)
        style.paragraph_format.space_after = Pt(4)

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.keep_with_next = False

    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        _set_font(run, "Arial", 8)
        run.font.color.rgb = RGBColor(100, 100, 100)
    _add_page_number(section.footer.paragraphs[0])

    doc.core_properties.title = title
    doc.core_properties.subject = "COCO publication validation"
    doc.core_properties.author = "Acycle COCO publication validation workflow"
    doc.core_properties.keywords = "cyclostratigraphy, COCO, Blocked cvCOCO, AR(1), periodogram"


def _add_title_page(doc: Any, title: str, subtitle: str, metadata: Sequence[tuple[str, str]]) -> None:
    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(42)
    paragraph.add_run(title)
    if subtitle:
        subtitle_paragraph = doc.add_paragraph(style="Subtitle")
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_paragraph.add_run(subtitle)
    doc.add_paragraph("")
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in metadata:
        cells = table.add_row().cells
        cells[0].width = Mm(46)
        cells[1].width = Mm(124)
        cells[0].text = label
        cells[1].text = value
        _shade_cell(cells[0], "D9EAF7")
        for cell in cells:
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                _set_font(run, "Times New Roman", 9.5, cell is cells[0])
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(16)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Primary inference: bidirectional held-out Blocked cvCOCO. "
        "Adaptive COCO is reported as exploratory."
    )
    _set_font(run, "Times New Roman", 10, True)
    doc.add_page_break()


def _add_key_value_table(
    doc: Any,
    rows: Sequence[tuple[str, str]],
    headers: tuple[str, str] = ("Item", "Value"),
    compact: bool = False,
) -> Any:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Mm(68)
    table.columns[1].width = Mm(112)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text, header_cells[1].text = headers
    _set_repeat_table_header(table.rows[0])
    for cell in header_cells:
        _shade_cell(cell, "1F4E79")
        for run in cell.paragraphs[0].runs:
            _set_font(run, "Arial", 8.5 if compact else 9, True)
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row_index, (label, value) in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = _text(label)
        cells[1].text = _text(value)
        if row_index % 2:
            for cell in cells:
                _shade_cell(cell, "F2F6FA")
        for cell in cells:
            _set_cell_margins(cell, 35 if compact else 55, 65, 35 if compact else 55, 65)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                _set_font(run, "Times New Roman", 8 if compact else 9)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    return table


def _add_bullets(doc: Any, values: Iterable[Any]) -> None:
    for value in values:
        text = _text(value)
        if not text:
            continue
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(text)


def _add_methods(doc: Any, manifest: Mapping[str, Any]) -> None:
    doc.add_heading("Methods and decision rules", level=1)
    for heading, body in STANDARD_METHODS:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)
    custom = manifest.get("methods")
    if custom:
        doc.add_heading("Run-specific methods", level=2)
        if isinstance(custom, Mapping):
            for heading, body in custom.items():
                paragraph = doc.add_paragraph()
                lead = paragraph.add_run(f"{_text(heading)}. ")
                lead.bold = True
                paragraph.add_run(_text(body))
        else:
            for item in _as_list(custom):
                if isinstance(item, Mapping):
                    heading = _text(
                        item.get("title") or item.get("method") or item.get("name"),
                        "Run-specific method",
                    )
                    body = _text(
                        item.get("description")
                        or item.get("details")
                        or item.get("role")
                        or item.get("value")
                    )
                    paragraph = doc.add_paragraph()
                    lead = paragraph.add_run(f"{heading}. ")
                    lead.bold = True
                    paragraph.add_run(body or "See the case parameter audit and run log.")
                else:
                    _add_bullets(doc, [item])


def _case_role(case: Mapping[str, Any]) -> str:
    explicit = case.get("design_role") or case.get("kind") or case.get("role")
    if explicit:
        return _text(explicit)
    case_id = _text(case.get("id")).casefold()
    searchable = (
        _text(case.get("title")) + " " + _text(case.get("expected_rate"))
    ).casefold()
    if "negative control" in searchable or "pure noise" in searchable:
        return "Negative control"
    if (
        "4to6" in case_id
        or "4_to_6" in case_id
        or "4-to-6" in searchable
        or "4 to 6" in searchable
    ) and ("la04" in case_id or "la2004" in searchable):
        return "Variable-rate synthetic stress test"
    if "la04" in case_id or "la2004" in searchable or "1e1t1p" in searchable:
        return "Positive synthetic control"
    if "pure astronomical signal" in searchable or "synthetic signal" in searchable:
        return "Positive synthetic control"
    return "Observed paleoclimate record"


def _expected_rate(case: Mapping[str, Any]) -> str:
    value = case.get("expected_rate_cm_per_kyr", case.get("expected_rate", "Not specified"))
    result = _text(value, "Not specified")
    lower = result.casefold()
    non_rate_label = (
        lower in {"none", "not specified", "n/a"}
        or "no physical" in lower
        or "negative control" in lower
    )
    if result and "cm" not in lower and not non_rate_label:
        result += " cm/kyr"
    return result


def _case_decisions(case: Mapping[str, Any]) -> dict[str, str]:
    rows = case.get("_summary", [])
    cv_class = _find_metric(rows, ("classification",), ("cv", "confirm"))
    if not cv_class:
        cv_class = _find_metric(rows, ("final", "confirmatory", "decision"))
    adaptive_class = _find_metric(rows, ("classification",), ("adaptive",))
    if not adaptive_class:
        adaptive_class = _find_metric(rows, ("final", "exploratory", "decision"))
    robust = _find_metric(rows, ("robust",))
    psym = _find_metric(rows, ("sym",))
    adaptive_p = _find_metric(rows, ("minimum", "global", "p"), ("adaptive",))
    if not adaptive_p:
        adaptive_p = _find_metric(rows, ("adaptive", "global", "p"))
    return {
        "cv": cv_class or "See case conclusion",
        "adaptive": adaptive_class or "See case conclusion",
        "p_robust": robust or "Not available",
        "p_sym": psym or "Not available",
        "adaptive_p": adaptive_p or "Not available",
    }


def _short(value: str, limit: int = 90) -> str:
    value = re.sub(r"\s+", " ", value).strip()
    return value if len(value) <= limit else value[: limit - 1].rstrip() + "…"


def _add_executive_summary(doc: Any, cases: Sequence[Mapping[str, Any]]) -> None:
    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        "The confirmatory interpretation is governed by Blocked cvCOCO and its "
        "preregistered p_robust < 0.05 rule. The Adaptive COCO result is retained "
        "as an exploratory diagnostic and should not replace the held-out result."
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Mm(37), Mm(27), Mm(32), Mm(42), Mm(42))
    headings = ("Data set", "Role", "Expected rate", "Blocked cvCOCO", "Adaptive COCO")
    for cell, heading, width in zip(table.rows[0].cells, headings, widths):
        cell.text = heading
        cell.width = width
        _shade_cell(cell, "1F4E79")
        for run in cell.paragraphs[0].runs:
            _set_font(run, "Arial", 7.5, True)
            run.font.color.rgb = RGBColor(255, 255, 255)
    _set_repeat_table_header(table.rows[0])
    for index, case in enumerate(cases):
        decisions = _case_decisions(case)
        values = (
            _text(case.get("title", case.get("id", ""))),
            _case_role(case),
            _expected_rate(case),
            decisions["cv"],
            decisions["adaptive"],
        )
        cells = table.add_row().cells
        for cell, value, width in zip(cells, values, widths):
            cell.text = _short(value)
            cell.width = width
            _set_cell_margins(cell, 40, 35, 40, 35)
            for run in cell.paragraphs[0].runs:
                _set_font(run, "Times New Roman", 7)
        if index % 2:
            for cell in cells:
                _shade_cell(cell, "F2F6FA")


def _case_metadata(case: Mapping[str, Any]) -> list[tuple[str, str]]:
    input_file = _text(case.get("input_file"), "Not recorded")
    age = _text(case.get("age_ma"), "Not specified")
    if age.casefold() != "not specified":
        age += " Ma"
    rows = [
        ("Case ID", _text(case.get("id"))),
        ("Case status", _text(case.get("status"), "Not recorded")),
        ("Design role", _case_role(case)),
        ("Input file", input_file),
        ("Input SHA-256", _text(case.get("input_sha256"), "Not recorded")),
        ("Middle age", age),
        ("Expected sedimentation rate", _expected_rate(case)),
    ]
    for key, label in (
        ("sample_count", "Input point count"),
        ("depth_range_m", "Depth range"),
        ("sampling_interval_m", "Sampling interval"),
    ):
        if case.get(key) not in (None, ""):
            rows.append((label, _text(case.get(key))))
    return rows


def _compact_parameter_rows(rows: Sequence[Mapping[str, str]]) -> list[tuple[str, str]]:
    priority = (
        "age", "sedimentation", "rate", "step", "maximum", "frequency",
        "nsim", "montecarlo", "seed", "ar1", "rednoise", "rho", "pad",
        "nfft", "slice", "sampling", "interpol",
    )
    selected: list[tuple[str, str]] = []
    for row in rows:
        metric = _text(row.get("Metric"))
        key = _normalized_key(metric)
        if any(token in key for token in priority):
            method = _text(row.get("Method"))
            label = f"{method}: {metric}" if method else metric
            selected.append((label, _text(row.get("Value"))))
    return selected[:24]


def _summary_rows(rows: Sequence[Mapping[str, str]], compact: bool) -> list[tuple[str, str]]:
    if not compact:
        return [
            (
                f"{_text(row.get('Method'))}: {_text(row.get('Metric'))}"
                if _text(row.get("Method")) else _text(row.get("Metric")),
                _text(row.get("Value")),
            )
            for row in rows
        ]
    tokens = (
        "classification", "decision", "conclusion", "best", "minimumglobal",
        "robust", "psym", "confidence", "participating", "allnine", "rho",
    )
    selected: list[tuple[str, str]] = []
    for row in rows:
        metric = _text(row.get("Metric"))
        if any(token in _normalized_key(metric) for token in tokens):
            method = _text(row.get("Method"))
            selected.append((f"{method}: {metric}" if method else metric, _text(row.get("Value"))))
    return selected[:32]


def _image_size_mm(path: Path, width_cap: float, height_cap: float) -> tuple[float, float | None]:
    if Image is None:
        return width_cap, None
    try:
        with Image.open(path) as image:
            pixel_width, pixel_height = image.size
    except Exception:
        return width_cap, None
    if pixel_width <= 0 or pixel_height <= 0:
        return width_cap, None
    ratio = pixel_height / pixel_width
    width = min(width_cap, height_cap / ratio)
    return width, width * ratio


def _add_picture_to_paragraph(paragraph: Any, figure: Mapping[str, Any], width_cap: float, height_cap: float) -> None:
    width, height = _image_size_mm(Path(figure["path"]), width_cap, height_cap)
    run = paragraph.add_run()
    if height is None:
        inline = run.add_picture(str(figure["path"]), width=Mm(width))
    else:
        inline = run.add_picture(str(figure["path"]), width=Mm(width), height=Mm(height))
    try:
        inline._inline.docPr.set("descr", _text(figure.get("caption")))
    except Exception:
        pass


def _add_caption(container: Any, number: int, figure: Mapping[str, Any]) -> None:
    paragraph = container.add_paragraph(style="Figure Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lead = paragraph.add_run(f"Figure {number}. ")
    lead.bold = True
    paragraph.add_run(_text(figure.get("caption")))
    method = _text(figure.get("method"))
    if method:
        paragraph.add_run(f" [{method}]")


def _add_figures(doc: Any, figures: Sequence[Mapping[str, Any]], start_number: int = 1) -> int:
    number = start_number
    index = 0
    while index < len(figures):
        figure = figures[index]
        if figure.get("width") == "half":
            pair = [figure]
            if index + 1 < len(figures) and figures[index + 1].get("width") == "half":
                pair.append(figures[index + 1])
            table = doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.columns[0].width = Mm(90)
            table.columns[1].width = Mm(90)
            for column in range(2):
                for row in range(2):
                    _set_cell_margins(table.cell(row, column), 0, 45, 0, 45)
            for column, item in enumerate(pair):
                image_paragraph = table.cell(0, column).paragraphs[0]
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_picture_to_paragraph(
                    image_paragraph, item, HALF_FIGURE_WIDTH_MM, MAX_HALF_FIGURE_HEIGHT_MM
                )
                caption_cell = table.cell(1, column)
                caption_cell.paragraphs[0]._element.getparent().remove(
                    caption_cell.paragraphs[0]._element
                )
                _add_caption(caption_cell, number, item)
                number += 1
            if len(pair) == 1:
                table.cell(0, 1).text = ""
                table.cell(1, 1).text = ""
            index += len(pair)
        else:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            _add_picture_to_paragraph(
                paragraph, figure, FULL_FIGURE_WIDTH_MM, MAX_FULL_FIGURE_HEIGHT_MM
            )
            _add_caption(doc, number, figure)
            number += 1
            index += 1
    return number


def _add_conclusion(doc: Any, case: Mapping[str, Any]) -> None:
    conclusion = _text(case.get("_conclusion"))
    if not conclusion:
        conclusion = _find_metric(case.get("_summary", []), ("conclusion", "report"))
    if not conclusion:
        conclusion = "No conclusion text was available; consult summary.csv."
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Mm(5)
    paragraph.paragraph_format.right_indent = Mm(5)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(conclusion)
    _set_font(run, "Times New Roman", 10, True)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "EAF2F8")
    paragraph._p.get_or_add_pPr().append(shading)


def _supporting_files(case: Mapping[str, Any]) -> list[tuple[str, str]]:
    root = Path(case.get("_root", case["_case_dir"]))
    case_dir = Path(case["_case_dir"])
    rows = [
        ("Case manifest", _text(case.get("_case_manifest"))),
        ("Parameters CSV", _text(case.get("_parameter_path"))),
        ("Summary CSV", _text(case.get("_summary_path"))),
        ("Conclusion text", _text(case.get("_conclusion_path"))),
    ]
    for filename, label in (
        ("parameters.mat", "Parameter structure and orbital periods"),
        ("preprocessed_input.csv", "Regularized analysis input"),
        ("case_results_index.mat", "Case result index"),
        ("run.log", "Case run log"),
    ):
        path = case_dir / filename
        if path.exists():
            rows.append((label, str(path)))
    for method_index, method in enumerate(_as_list(case.get("methods")), 1):
        if not isinstance(method, Mapping):
            continue
        method_name = _text(method.get("method"), f"Method {method_index}")
        role = _text(method.get("role"))
        label_prefix = f"{method_name} ({role})" if role else method_name
        for field, label in (
            ("result_file", "MATLAB result"),
            ("workbook", "result workbook"),
            ("conclusion_file", "method conclusion"),
        ):
            raw_path = method.get(field)
            if not raw_path:
                continue
            path = _resolve_file(root, case_dir, raw_path)
            if path.exists():
                rows.append((f"{label_prefix} {label}", str(path)))
    return rows


def _add_case_section(
    doc: Any,
    case: Mapping[str, Any],
    figure_number: int,
    compact: bool,
    page_break: bool,
) -> int:
    if page_break:
        doc.add_page_break()
    doc.add_heading(_text(case.get("title", case.get("id"))), level=1)
    _add_key_value_table(doc, _case_metadata(case), headers=("Case attribute", "Recorded value"), compact=True)

    doc.add_heading("Conclusion", level=2)
    _add_conclusion(doc, case)

    decisions = _case_decisions(case)
    _add_key_value_table(
        doc,
        [
            ("Blocked cvCOCO classification", decisions["cv"]),
            ("p_robust = max(p_A, p_B)", decisions["p_robust"]),
            ("Secondary p_sym", decisions["p_sym"]),
            ("Adaptive COCO classification", decisions["adaptive"]),
            ("Adaptive minimum global p", decisions["adaptive_p"]),
        ],
        headers=("Decision metric", "Result"),
        compact=True,
    )

    parameters = case.get("_parameters", [])
    if parameters:
        doc.add_heading("Analysis parameters", level=2)
        rows = _compact_parameter_rows(parameters) if compact else _summary_rows(parameters, False)
        _add_key_value_table(doc, rows, headers=("Parameter", "Value"), compact=True)
        if compact and len(rows) < len(parameters):
            doc.add_paragraph(
                f"The complete {len(parameters)}-row parameter audit is stored in parameters.csv."
            )

    summary = case.get("_summary", [])
    if summary:
        doc.add_heading("Numerical result summary", level=2)
        rows = _summary_rows(summary, compact)
        _add_key_value_table(doc, rows, headers=("Metric", "Value"), compact=True)
        if compact and len(rows) < len(summary):
            doc.add_paragraph(
                f"The complete {len(summary)}-row decision audit is stored in summary.csv."
            )

    notes = _as_list(case.get("notes")) + _as_list(case.get("quality_notes"))
    if notes:
        doc.add_heading("Case-specific interpretation and quality notes", level=2)
        _add_bullets(doc, notes)

    figures = case.get("_figures", [])
    if figures:
        doc.add_heading("Figures", level=2)
        figure_number = _add_figures(doc, figures, figure_number)
    else:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("No embeddable PNG figure was supplied for this case.")
        run.italic = True

    if not compact:
        doc.add_heading("Supporting audit files", level=2)
        _add_key_value_table(
            doc, _supporting_files(case), headers=("Artifact", "Path"), compact=True
        )
        case_warnings = _as_list(case.get("_warnings"))
        if case_warnings:
            doc.add_heading("Nonfatal report-input warnings", level=2)
            _add_bullets(doc, case_warnings)
    return figure_number


def _all_limitations(manifest: Mapping[str, Any]) -> list[str]:
    limitations = list(STANDARD_LIMITATIONS)
    for item in _as_list(manifest.get("limitations")):
        text = _text(item)
        if text and text not in limitations:
            limitations.append(text)
    return limitations


def _generation_time(manifest: Mapping[str, Any]) -> str:
    value = (
        manifest.get("generated_at")
        or manifest.get("run_date")
        or manifest.get("updated_at")
        or manifest.get("created_at")
    )
    if value:
        return _text(value)
    return _datetime.datetime.now(_datetime.timezone.utc).astimezone().isoformat(timespec="seconds")


def _software_text(manifest: Mapping[str, Any]) -> str:
    software = manifest.get("software")
    if software:
        return _text(software)
    pieces = []
    git_revision = manifest.get("git_commit") or manifest.get("git_head")
    if git_revision:
        pieces.append(f"Git: {_text(git_revision)}")
    for key, label in (
        ("matlab_version", "MATLAB"),
        ("acycle_version", "Acycle"),
    ):
        if manifest.get(key):
            pieces.append(f"{label}: {_text(manifest[key])}")
    return "; ".join(pieces) or "See manifest.json and run logs"


def build_combined_document(
    root: Path,
    manifest: Mapping[str, Any],
    cases: Sequence[Mapping[str, Any]],
    warnings: WarningLog,
) -> Any:
    _require_docx()
    title = _text(
        manifest.get("report_title") or manifest.get("title"),
        "COCO publication validation report",
    )
    subtitle = _text(
        manifest.get("report_subtitle"),
        "Confirmatory bidirectional held-out Blocked cvCOCO and exploratory Adaptive COCO",
    )
    doc = Document()
    _configure_document(doc, title)
    _add_title_page(
        doc,
        title,
        subtitle,
        (
            ("Result root", str(root)),
            ("Cases", str(len(cases))),
            ("Suite status", _text(manifest.get("status"), "Not recorded")),
            ("Generated", _generation_time(manifest)),
            ("Software provenance", _software_text(manifest)),
            ("Report builder", f"build_coco_publication_reports.py {SCRIPT_VERSION}"),
        ),
    )
    _add_executive_summary(doc, cases)
    _add_methods(doc, manifest)
    figure_number = 1
    for index, case in enumerate(cases):
        figure_number = _add_case_section(
            doc, case, figure_number, compact=True, page_break=True
        )
    doc.add_page_break()
    doc.add_heading("Interpretive limitations", level=1)
    _add_bullets(doc, _all_limitations(manifest))
    doc.add_heading("Reproducibility and report scope", level=1)
    doc.add_paragraph(
        "This Word document is a human-readable rendering of the immutable "
        "manifest, CSV summaries, conclusion texts, and PNG figures. Numerical "
        "claims should be audited against the saved MATLAB result structures, "
        "workbooks, run logs, file hashes, and source revision recorded in the "
        "result directory. The report builder does not recompute statistics or "
        "alter conclusions."
    )
    if warnings.messages:
        doc.add_heading("Nonfatal report-generation warnings", level=2)
        _add_bullets(doc, warnings.messages)
    return doc


def build_case_document(
    root: Path,
    manifest: Mapping[str, Any],
    case: Mapping[str, Any],
) -> Any:
    _require_docx()
    title = f"COCO case report: {_text(case.get('title', case.get('id')))}"
    doc = Document()
    _configure_document(doc, title)
    _add_title_page(
        doc,
        title,
        "Confirmatory Blocked cvCOCO and exploratory Adaptive COCO",
        (
            ("Case ID", _text(case.get("id"))),
            ("Result directory", str(case.get("_case_dir"))),
            ("Generated", _generation_time(manifest)),
            ("Software provenance", _software_text(manifest)),
            ("Report builder", f"build_coco_publication_reports.py {SCRIPT_VERSION}"),
        ),
    )
    _add_case_section(doc, case, figure_number=1, compact=False, page_break=False)
    doc.add_page_break()
    _add_methods(doc, manifest)
    doc.add_heading("Interpretive limitations", level=1)
    _add_bullets(doc, _all_limitations(manifest))
    return doc


def _atomic_save(doc: Any, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    handle, temporary_name = tempfile.mkstemp(
        prefix=f".{destination.stem}.", suffix=".tmp.docx", dir=destination.parent
    )
    os.close(handle)
    temporary = Path(temporary_name)
    try:
        doc.save(str(temporary))
        os.replace(temporary, destination)
    except Exception:
        try:
            temporary.unlink()
        except OSError:
            pass
        raise


def _safe_case_id(case: Mapping[str, Any], index: int) -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", _text(case.get("id"))).strip("._")
    return value or f"case_{index:02d}"


def parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build A4 combined and per-case Word reports for a COCO suite."
    )
    parser.add_argument(
        "source",
        type=Path,
        help="Result root, manifest.json, or CSV manifest.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=None,
        help=(
            "Optional report destination. Without it, the combined report is "
            "written to the result root and each case report to its case directory."
        ),
    )
    parser.add_argument(
        "--combined-name",
        default="COCO_publication_validation_report.docx",
        help="Filename for the combined report.",
    )
    parser.add_argument(
        "--case-name",
        default="case_report.docx",
        help="Filename for each case report.",
    )
    parser.add_argument(
        "--no-case-reports",
        action="store_true",
        help="Build only the combined report.",
    )
    parser.add_argument("--version", action="version", version=SCRIPT_VERSION)
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> int:
    args = parse_args(argv)
    warnings = WarningLog()
    try:
        root, manifest, cases = load_suite(args.source, warnings)
        destination_root = args.output_dir.expanduser().resolve() if args.output_dir else root
        combined_path = destination_root / args.combined_name
        combined = build_combined_document(root, manifest, cases, warnings)
        _atomic_save(combined, combined_path)
        print(f"Wrote combined report: {combined_path}")

        if not args.no_case_reports:
            for index, case in enumerate(cases, 1):
                if args.output_dir:
                    case_destination = destination_root / _safe_case_id(case, index) / args.case_name
                else:
                    case_destination = Path(case["_case_dir"]) / args.case_name
                case_doc = build_case_document(root, manifest, case)
                _atomic_save(case_doc, case_destination)
                print(f"Wrote case report: {case_destination}")
    except ReportInputError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    except Exception as exc:
        print(f"ERROR: report generation failed: {exc}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
